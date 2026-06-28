"""Word 课表导出

从模板 .docx 生成课表，每页一份完整课表（标题+副标题+表格+落款+日期）。
"""

import io
import zipfile
from collections import defaultdict
from copy import deepcopy
from pathlib import Path

from django.http import HttpResponse
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from rest_framework import status
from rest_framework.decorators import api_view
from rest_framework.response import Response

from core.models import SchoolClass, Subject, Teacher, TravelGroup, SchedulerSettings
from .models import ScheduleEntry, ScheduleResult
from .time_slots import DAYS

# 模板路径 — 相对 backend/ 运行目录
TEMPLATE_PATH = Path(__file__).resolve().parent.parent / 'assets' / '个人课程表模版.docx'

# 节次 → 表格列索引 (0-based)
# col 0=星期标签, col 1=第1节, col 2=第2节, col 3=阳光大课间(固定),
# col 4=第3节, col 5=眼保健操(固定), col 6=第4节, col 7=午休(固定),
# col 8=第5节, col 9=第6节, col 10=阳光体育(固定)
PERIOD_COLUMN = {0: 1, 1: 2, 2: 4, 3: 6, 4: 8, 5: 9}

# 数据行起始行号：模板 Row 0/1 是表头，Row 2~6 是周一到周五
DATA_ROW_OFFSET = 2

# ── 模板结构索引 (body children) ──
# 这些索引由 analyse_template.py 输出，模板改动后需重新确认
IDX_TITLE = 0       # "课  程  表"
IDX_SUBTITLE = 1    # "某某某学校2025年下学期  ...  某某某"
IDX_TABLE = 2       # 课程表表格
IDX_SPACER_START = 3  # 表格后的空段落（用于撑开版式）
IDX_SPACER_END = 26
IDX_FOOTER = 27     # "教导处"
IDX_DATE = 28        # "2025年8月"
# IDX_SECTPR = 29    # 页面设置，整个文档只有一个


# ── helpers ────────────────────────────────────────────────────

def _make_page_break():
    """创建分页符 <w:p><w:r><w:br w:type="page"/></w:r></w:p>"""
    p = etree.Element(qn('w:p'))
    r = etree.SubElement(p, qn('w:r'))
    etree.SubElement(r, qn('w:br'), {qn('w:type'): 'page'})
    return p


# 页面度量常数（twips），用于让段落左右边界与表格对齐
# 模板: A4 横向 16838×11906, 左右页边距 1440, 表格宽 12678 居中
PAGE_WIDTH = 16838
PAGE_MARGIN = 1440
TABLE_WIDTH = 12678
# 表格相对文本区的左右缩进值
TABLE_LEFT_INDENT = (PAGE_WIDTH - TABLE_WIDTH) // 2 - PAGE_MARGIN  # = 640
TABLE_RIGHT_INDENT = TABLE_LEFT_INDENT
# 头像/落款/日期相对表格边缘的微调距离（约一个中文字的间距）
INNER_INDENT = 300


def _ensure_pPr(para_element):
    """获取或创建段落属性元素"""
    pPr = para_element.find(qn('w:pPr'))
    if pPr is None:
        pPr = etree.SubElement(para_element, qn('w:pPr'))
    return pPr


def _apply_indent(para_element, *, left: int = None, right: int = None,
                  alignment: str = None):
    """为段落设置左右缩进（覆盖旧值），可选对齐方式"""
    pPr = _ensure_pPr(para_element)
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = etree.Element(qn('w:ind'))
        pPr.insert(0, ind)
    if left is not None:
        ind.set(qn('w:left'), str(left))
    if right is not None:
        ind.set(qn('w:right'), str(right))
    if alignment:
        jc = pPr.find(qn('w:jc'))
        if jc is None:
            jc = etree.Element(qn('w:jc'))
            pPr.insert(0, jc)
        jc.set(qn('w:val'), alignment)


def _clear_paragraph_runs(para_element):
    """移除段落中所有 <w:r>"""
    for r in list(para_element.findall(qn('w:r'))):
        para_element.remove(r)


def _add_paragraph_run(para_element, text: str, *, bold: bool = True, size_pt: int = 0):
    """向段落添加一个 run（可选字号，单位 pt）"""
    if not text:
        return
    r = etree.SubElement(para_element, qn('w:r'))
    rPr = etree.SubElement(r, qn('w:rPr'))
    etree.SubElement(rPr, qn('w:b'))
    etree.SubElement(rPr, qn('w:bCs'))
    if size_pt > 0:
        sz_str = str(size_pt * 2)  # half-points
        sz = etree.SubElement(rPr, qn('w:sz'))
        sz.set(qn('w:val'), sz_str)
        szCs = etree.SubElement(rPr, qn('w:szCs'))
        szCs.set(qn('w:val'), sz_str)
    t = etree.SubElement(r, qn('w:t'))
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


def _set_paragraph_text(para_element, new_text: str):
    """清空段落并写入纯文本（保留原属性）"""
    _clear_paragraph_runs(para_element)
    if new_text:
        _add_paragraph_run(para_element, new_text)


def _set_subtitle_text(para_element, school_semester: str, target_name: str):
    """副标题：学校+学期靠左（较表格左边缘略右缩），教师/班级名靠右（较表格右边缘略左缩）

    左右缩进 = 表格缩进 + INNER_INDENT，使文本从表格边缘向内收敛。
    右制表位 = TABLE_WIDTH - INNER_INDENT，让名字也与右边缘保持相同间距。
    """
    _clear_paragraph_runs(para_element)
    if not school_semester and not target_name:
        _set_paragraph_text(para_element, '')
        return

    pPr = _ensure_pPr(para_element)
    jc_old = pPr.find(qn('w:jc'))
    if jc_old is not None:
        pPr.remove(jc_old)
    # 清理模板残留的段落级 rPr
    old_rPr = pPr.find(qn('w:rPr'))
    if old_rPr is not None:
        pPr.remove(old_rPr)

    _apply_indent(para_element,
                  left=TABLE_LEFT_INDENT + INNER_INDENT,
                  right=TABLE_RIGHT_INDENT + INNER_INDENT,
                  alignment='left')

    # 右制表位：段落内可用宽度 = TABLE_WIDTH - INNER_INDENT
    tabs = pPr.find(qn('w:tabs'))
    if tabs is None:
        tabs = etree.SubElement(pPr, qn('w:tabs'))
    else:
        for t in list(tabs):
            tabs.remove(t)
    tab = etree.SubElement(tabs, qn('w:tab'))
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(TABLE_WIDTH - INNER_INDENT))

    if school_semester:
        _add_paragraph_run(para_element, school_semester, size_pt=16)

    r_tab = etree.SubElement(para_element, qn('w:r'))
    etree.SubElement(r_tab, qn('w:tab'))

    if target_name:
        _add_paragraph_run(para_element, target_name, size_pt=16)


def _set_footer_center(para_element, text: str):
    """落款或日期：在行右侧区域居中对齐，段间距紧凑"""
    _clear_paragraph_runs(para_element)
    pPr = _ensure_pPr(para_element)
    # 清理模板残留的段落级 rPr（旧字号会干扰）
    old_rPr = pPr.find(qn('w:rPr'))
    if old_rPr is not None:
        pPr.remove(old_rPr)
    # 左缩进推到表格右 1/4 区域再居中；右缩进保持与表格右边缘的间距
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = etree.Element(qn('w:ind'))
        pPr.insert(0, ind)
    ind.set(qn('w:left'), str(TABLE_LEFT_INDENT + TABLE_WIDTH * 3 // 4))
    ind.set(qn('w:right'), str(TABLE_RIGHT_INDENT + INNER_INDENT))
    # 对齐
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = etree.Element(qn('w:jc'))
        pPr.insert(0, jc)
    jc.set(qn('w:val'), 'center')
    # 段间距：单倍行距，段前段后为 0
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = etree.Element(qn('w:spacing'))
        pPr.insert(0, spacing)
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'auto')
    # 关闭 snapToGrid，否则 docGrid(312) 会撑大段落间距
    snap = pPr.find(qn('w:snapToGrid'))
    if snap is None:
        snap = etree.Element(qn('w:snapToGrid'))
        pPr.append(snap)
    snap.set(qn('w:val'), '0')
    if text:
        _add_paragraph_run(para_element, text, size_pt=14)


def _get_table_cells(tbl_element):
    """返回表格中所有单元格 XML 元素，按 (row_idx, col_idx) 组织"""
    rows_xml = tbl_element.findall(qn('w:tr'))
    cells = []
    for row in rows_xml:
        row_cells = row.findall(qn('w:tc'))
        cells.append(row_cells)
    return cells


def _clear_cell_paragraphs(cell_element):
    """移除单元格内所有 <w:p>"""
    for p in list(cell_element.findall(qn('w:p'))):
        cell_element.remove(p)


def _add_cell_paragraph(cell_element, text: str, *, bold: bool = False,
                        size_pt: int = 0, alignment: str = 'center'):
    """向单元格添加一个段落"""
    p = etree.SubElement(cell_element, qn('w:p'))
    pPr = etree.SubElement(p, qn('w:pPr'))
    jc = etree.SubElement(pPr, qn('w:jc'))
    jc.set(qn('w:val'), alignment)
    r = etree.SubElement(p, qn('w:r'))
    rPr = etree.SubElement(r, qn('w:rPr'))
    if bold:
        etree.SubElement(rPr, qn('w:b'))
        etree.SubElement(rPr, qn('w:bCs'))
    if size_pt > 0:
        sz_val = str(size_pt * 2)  # half-points
        sz = etree.SubElement(rPr, qn('w:sz'))
        sz.set(qn('w:val'), sz_val)
        szCs = etree.SubElement(rPr, qn('w:szCs'))
        szCs.set(qn('w:val'), sz_val)
    t = etree.SubElement(r, qn('w:t'))
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


# ── 数据准备 ────────────────────────────────────────────────────

def _is_class_meeting(entry_or_subject):
    """判断一个 entry 或 subject 是否为班会课"""
    settings = SchedulerSettings.get_settings()
    meeting_name = settings.class_meeting_name
    # ScheduleEntry 有 subject 外键
    if hasattr(entry_or_subject, 'subject'):
        subj = entry_or_subject.subject
        return bool(subj and subj.name == meeting_name)
    # Subject 实例
    return bool(entry_or_subject and entry_or_subject.name == meeting_name)


def _is_combined(entry_or_subject):
    """判断一个 entry 或 subject 是否为校本课程(合班课)"""
    if hasattr(entry_or_subject, 'subject'):
        subj = entry_or_subject.subject
        return bool(subj and subj.is_combined_class)
    return bool(entry_or_subject and entry_or_subject.is_combined_class)


def _build_entry_maps(result_id, view_type, targets):
    """构建 {target_id: [(day, period, cell_line1, cell_line2)]} 映射

    返回值中每个条目是 (day, period, line1, line2) —— 已经算好显示文本。
    """
    entries_qs = ScheduleEntry.objects.filter(
        result_id=result_id
    ).select_related('subject', 'teacher', 'school_class')

    # 获取班会/校本课程标识
    settings = SchedulerSettings.get_settings()
    meeting_name = settings.class_meeting_name

    if view_type == 'class':
        # ── 班级视图 ──
        groups = defaultdict(list)
        for e in entries_qs:
            groups[e.school_class_id].append(e)

        maps = {}
        for t in targets:
            entry_list = groups.get(t.id, [])
            cell_data = []
            for e in entry_list:
                subject = e.subject.name if e.subject_id else ''
                teacher = e.teacher.name if e.teacher_id else ''
                cell_data.append((e.day, e.period, subject, teacher))
            maps[t.id] = cell_data
        return maps

    else:
        # ── 教师视图 ──
        groups = defaultdict(list)
        for e in entries_qs:
            if e.teacher_id:
                groups[e.teacher_id].append(e)

        # 补充校本课程条目
        try:
            result = ScheduleResult.objects.get(pk=result_id)
            combined_assignments = result.combined_class_assignments or {}
        except ScheduleResult.DoesNotExist:
            combined_assignments = {}

        if combined_assignments:
            combined_slots = settings.get_combined_class_slots_list()
            # 构建 teacher_name → assigned_days 映射
            teacher_days = {}  # teacher_id → [day, ...]
            teacher_name_to_id = {t.name: t.id for t in targets}
            for group_name, day_data in combined_assignments.items():
                if not isinstance(day_data, dict):
                    continue
                for day_label, names in day_data.items():
                    day = 1 if '周二' in day_label else 3 if '周四' in day_label else None
                    if day is None:
                        continue
                    for t_name in names:
                        tid = teacher_name_to_id.get(t_name)
                        if tid:
                            teacher_days.setdefault(tid, []).append(day)

            for tid, days in teacher_days.items():
                for day in days:
                    for slot_day, slot_period in combined_slots:
                        if slot_day == day:
                            groups[tid].append(_FakeEntry(
                                day=slot_day, period=slot_period,
                                line1='校本课程', line2='',
                            ))

        # 预计算同年级同科目的班级区分后缀
        teacher_grade_subj_classes = defaultdict(lambda: defaultdict(set))
        for tid, entries in groups.items():
            for e in entries:
                if isinstance(e, _FakeEntry):
                    continue
                if e.school_class_id and e.subject_id and e.subject.name != meeting_name and not e.subject.is_combined_class:
                    key = (e.school_class.grade, e.subject.name)
                    teacher_grade_subj_classes[tid][key].add(e.school_class_id)

        maps = {}
        for t in targets:
            entry_list = groups.get(t.id, [])
            multi_class_keys = {
                k for k, cids in teacher_grade_subj_classes.get(t.id, {}).items()
                if len(cids) > 1
            }
            # 对 multi-class 的每组，给班级编号
            class_indices = {}
            for key in multi_class_keys:
                sorted_ids = sorted(teacher_grade_subj_classes[t.id][key])
                class_indices.update({cid: i + 1 for i, cid in enumerate(sorted_ids)})

            cell_data = []
            for e in entry_list:
                if hasattr(e, 'line1'):
                    # 校本/班会 的 FakeEntry
                    cell_data.append((e.day, e.period, e.line1, e.line2))
                elif _is_class_meeting(e):
                    cell_data.append((e.day, e.period, e.subject.name if e.subject_id else '班会', ''))
                elif _is_combined(e):
                    cell_data.append((e.day, e.period, '校本课程', ''))
                else:
                    # 普通课程
                    grade = e.school_class.grade if e.school_class_id else ''
                    subject_name = e.subject.name if e.subject_id else ''
                    line1 = f'{grade}{subject_name}'
                    # 同年级多班时加括号
                    cid = e.school_class_id
                    if cid in class_indices:
                        line1 = f'{line1}({class_indices[cid]})'
                    cell_data.append((e.day, e.period, line1, ''))
            maps[t.id] = cell_data
        return maps


class _FakeEntry:
    """用于校本课程等非 ScheduleEntry 的占位条目"""
    __slots__ = ('day', 'period', 'line1', 'line2')
    def __init__(self, day, period, line1, line2):
        self.day = day
        self.period = period
        self.line1 = line1
        self.line2 = line2


# ── 表格填充 & 页面克隆 ──────────────────────────────────────────

def _fill_table(tbl_element, cell_data):
    """用已算好的 cell_data 填充一个表格的课程格

    cell_data: [(day, period, line1, line2), ...]
    """
    cells = _get_table_cells(tbl_element)
    data_map = {(d, p): (l1, l2) for d, p, l1, l2 in cell_data}

    for day in range(5):
        row_idx = DATA_ROW_OFFSET + day
        if row_idx >= len(cells):
            continue

        for period in range(6):
            col_idx = PERIOD_COLUMN.get(period)
            if col_idx is None or col_idx >= len(cells[row_idx]):
                continue

            cell_el = cells[row_idx][col_idx]
            _clear_cell_paragraphs(cell_el)

            pair = data_map.get((day, period))
            if not pair:
                continue

            line1, line2 = pair
            _add_cell_paragraph(cell_el, line1, bold=True, size_pt=11)
            if line2:
                _add_cell_paragraph(cell_el, line2, bold=False, size_pt=8)


def _clone_page_elements(body, template_elements):
    """从模板元素列表中 deepcopy 一份"""
    return [deepcopy(el) for el in template_elements]


def _insert_sequence_after(prev_element, elements):
    """在 prev_element 之后依次插入 elements，返回最后插入的元素"""
    cur = prev_element
    for el in elements:
        cur.addnext(el)
        cur = el
    return cur


def _build_template_element_list(body):
    """抽取模板中组成一页的 XML 元素（标题 → 落款），不含 sectPr"""
    children = list(body)
    if len(children) <= IDX_DATE:
        raise RuntimeError('模板结构不符合预期，请检查模板文件')
    # 按顺序收集：标题 副标题 表格 空段落们 落款 日期
    elements = []
    elements.append(children[IDX_TITLE])
    elements.append(children[IDX_SUBTITLE])
    elements.append(children[IDX_TABLE])
    for i in range(IDX_SPACER_START, IDX_SPACER_END + 1):
        elements.append(children[i])
    # 落款和日期位置取决于空段落数，用实际内容段落来找
    # 从后往前找最后两个有实际内容的段落
    return elements


# ── main endpoint ──────────────────────────────────────────────

def _build_single_docx_bytes(result_id, view_type, targets, school_name, semester,
                             footer_text, date_text):
    """为一种导出类型生成一个 docx 文件（多页），返回 bytes。"""

    cell_data_maps = _build_entry_maps(result_id, view_type, targets)

    # 组装 (target, cell_data) 列表
    page_items = [(t, cell_data_maps.get(t.id, [])) for t in targets]

    return _generate_docx_from_pages(page_items, school_name, semester,
                                     footer_text, date_text)


def _generate_docx_from_pages(page_items, school_name, semester,
                              footer_text, date_text):
    """用模板生成一份多页 docx，每页一个 target 的课表。返回 bytes。"""

    doc = Document(str(TEMPLATE_PATH))
    body = doc.element.body

    # 压小下边距，防止落款/时间溢出到第二页（原 1236 → 450）
    sect_pr = body.find(qn('w:sectPr'))
    if sect_pr is not None:
        pg_mar = sect_pr.find(qn('w:pgMar'))
        if pg_mar is not None:
            pg_mar.set(qn('w:bottom'), '450')
    all_children = list(body)

    # 模板章节
    title_el = all_children[IDX_TITLE]
    subtitle_el = all_children[IDX_SUBTITLE]
    table_el = all_children[IDX_TABLE]
    spacer_els = []
    for i in range(IDX_SPACER_START, IDX_SPACER_END + 1):
        if i < len(all_children) - 1:
            spacer_els.append(all_children[i])
    footer_el = all_children[IDX_FOOTER]
    date_el = all_children[IDX_DATE]

    sect_pr = list(body)[-1]

    page_template = [title_el, subtitle_el, table_el] + spacer_els + [footer_el, date_el]

    # 填充第一页
    first_target, first_data = page_items[0]
    _set_subtitle_text(subtitle_el, f'{school_name}  {semester}', first_target.name)
    _fill_table(table_el, first_data)
    _set_footer_center(footer_el, footer_text)
    _set_footer_center(date_el, date_text)

    # 后续页面：克隆整页结构
    last_page_end = date_el
    for target, data in page_items[1:]:
        pb = _make_page_break()
        last_page_end.addnext(pb)
        cloned = _clone_page_elements(body, page_template)
        cl_title, cl_subtitle, cl_table = cloned[0], cloned[1], cloned[2]
        cl_footer, cl_date = cloned[-2], cloned[-1]

        _set_subtitle_text(cl_subtitle, f'{school_name}  {semester}', target.name)
        _fill_table(cl_table, data)
        _set_footer_center(cl_footer, footer_text)
        _set_footer_center(cl_date, date_text)

        last_page_end = _insert_sequence_after(pb, cloned)

    # 确保 sectPr 在最后
    if list(body)[-1] is not sect_pr:
        body.remove(sect_pr)
        body.append(sect_pr)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _make_para(text, *, bold=False, size_pt=14, alignment='left',
               left_indent=0, right_indent=0):
    """创建一个带样式的 w:p 元素。"""
    p = etree.Element(qn('w:p'))
    pPr = etree.SubElement(p, qn('w:pPr'))
    jc = etree.SubElement(pPr, qn('w:jc'))
    jc.set(qn('w:val'), alignment)
    if left_indent or right_indent:
        ind = etree.SubElement(pPr, qn('w:ind'))
        if left_indent:
            ind.set(qn('w:left'), str(left_indent))
        if right_indent:
            ind.set(qn('w:right'), str(right_indent))
    if text:
        r = etree.SubElement(p, qn('w:r'))
        rPr = etree.SubElement(r, qn('w:rPr'))
        if bold:
            etree.SubElement(rPr, qn('w:b'))
            etree.SubElement(rPr, qn('w:bCs'))
        if size_pt:
            sz_val = str(size_pt * 2)
            sz = etree.SubElement(rPr, qn('w:sz'))
            sz.set(qn('w:val'), sz_val)
            szCs = etree.SubElement(rPr, qn('w:szCs'))
            szCs.set(qn('w:val'), sz_val)
        t = etree.SubElement(r, qn('w:t'))
        t.text = text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return p


def _make_footer_para(text):
    """落款/日期段落：在行右侧区域居中对齐，与课表页面的落款风格一致。"""
    p = etree.Element(qn('w:p'))
    pPr = etree.SubElement(p, qn('w:pPr'))
    ind = etree.SubElement(pPr, qn('w:ind'))
    ind.set(qn('w:left'), str(TABLE_LEFT_INDENT + TABLE_WIDTH * 3 // 4))
    ind.set(qn('w:right'), str(TABLE_RIGHT_INDENT + INNER_INDENT))
    jc = etree.SubElement(pPr, qn('w:jc'))
    jc.set(qn('w:val'), 'center')
    spacing = etree.SubElement(pPr, qn('w:spacing'))
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'auto')
    snap = etree.SubElement(pPr, qn('w:snapToGrid'))
    snap.set(qn('w:val'), '0')
    if text:
        r = etree.SubElement(p, qn('w:r'))
        rPr = etree.SubElement(r, qn('w:rPr'))
        etree.SubElement(rPr, qn('w:b'))
        etree.SubElement(rPr, qn('w:bCs'))
        sz = etree.SubElement(rPr, qn('w:sz'))
        sz.set(qn('w:val'), '28')
        szCs = etree.SubElement(rPr, qn('w:szCs'))
        szCs.set(qn('w:val'), '28')
        t = etree.SubElement(r, qn('w:t'))
        t.text = text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return p


def _get_groups_data(result_id):
    """获取分组信息数据。返回 (combined_data, travel_data)。

    combined_data: {分组名: [教师名, ...]}
    travel_data:   [{'name': 分组名, 'day_off': 禁排日, 'teachers': [教师名, ...]}, ...]
    """
    try:
        result = ScheduleResult.objects.get(pk=result_id)
    except ScheduleResult.DoesNotExist:
        result = None

    combined_assignments = result.combined_class_assignments if result else {}

    # 合并周二/周四教师列表，不区分日期
    combined_data = {}
    for group_name, day_data in combined_assignments.items():
        teachers = []
        if isinstance(day_data, dict):
            for names in day_data.values():
                teachers.extend(names)
        elif isinstance(day_data, list):
            teachers = day_data
        combined_data[group_name] = teachers

    # 送教分组
    travel_groups = list(TravelGroup.objects.all().order_by('day_off'))
    teachers = list(Teacher.objects.all())
    teachers_by_group = {}
    for t in teachers:
        if t.travel_group_id:
            teachers_by_group.setdefault(t.travel_group_id, []).append(t.name)

    travel_data = []
    for tg in travel_groups:
        travel_data.append({
            'name': tg.name,
            'day_off': tg.get_day_off_display() or '',
            'teachers': teachers_by_group.get(tg.id, []),
        })

    return combined_data, travel_data


def _build_groups_page_elements(combined_data, travel_data, school_name, semester,
                                 footer_text, date_text):
    """构建分组信息页面的所有 w:p 元素。"""
    indent = TABLE_LEFT_INDENT + INNER_INDENT  # 与表格左边缘对齐，略向右微调

    elements = []
    # 标题
    elements.append(_make_para('课  程  表', bold=True, size_pt=22, alignment='center'))
    # 副标题
    elements.append(_make_para(f'{school_name}  {semester}  分组信息',
                               bold=True, size_pt=16, alignment='center'))
    # 标题与正文间距
    elements.append(_make_para('', size_pt=10))

    # ── 校本课程教师分组 ──
    elements.append(_make_para('校本课程教师分组',
                               bold=True, size_pt=14, alignment='left',
                               left_indent=indent))
    if combined_data:
        for group_name in sorted(combined_data.keys()):
            teacher_names = combined_data[group_name]
            t_str = '、'.join(teacher_names) if teacher_names else '（无）'
            elements.append(_make_para(f'{group_name}：{t_str}',
                                       size_pt=14, alignment='left',
                                       left_indent=indent + 400))
    else:
        elements.append(_make_para('（无）', size_pt=14, alignment='left',
                                   left_indent=indent + 400))

    # 两组之间的间距
    elements.append(_make_para('', size_pt=12))

    # ── 送教分组 ──
    elements.append(_make_para('送教分组',
                               bold=True, size_pt=14, alignment='left',
                               left_indent=indent))
    if travel_data:
        for tg in travel_data:
            label = f"{tg['name']}（{tg['day_off']}）" if tg['day_off'] else tg['name']
            t_str = '、'.join(tg['teachers']) if tg['teachers'] else '（无）'
            elements.append(_make_para(f'{label}：{t_str}',
                                       size_pt=14, alignment='left',
                                       left_indent=indent + 400))
    else:
        elements.append(_make_para('（无）', size_pt=14, alignment='left',
                                   left_indent=indent + 400))

    # 落款和日期（与课表页面同款样式）
    elements.append(_make_footer_para(footer_text))
    elements.append(_make_footer_para(date_text))

    return elements


def _generate_groups_docx_bytes(result_id, school_name, semester,
                                 footer_text, date_text):
    """生成仅包含分组信息的 docx 文件，返回 bytes。"""
    combined_data, travel_data = _get_groups_data(result_id)
    elements = _build_groups_page_elements(combined_data, travel_data,
                                           school_name, semester,
                                           footer_text, date_text)

    doc = Document(str(TEMPLATE_PATH))
    body = doc.element.body

    # 清空模板内容，只保留 sectPr（页面设置）
    children = list(body)
    sect_pr = children[-1]
    for child in children:
        body.remove(child)

    for el in elements:
        body.append(el)
    body.append(sect_pr)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _append_groups_to_body(body, combined_data, travel_data, school_name, semester,
                            footer_text, date_text, after_element):
    """在 after_element 之后追加分页符 + 分组信息页面元素。返回最后插入的元素。"""
    pb = _make_page_break()
    after_element.addnext(pb)

    elements = _build_groups_page_elements(combined_data, travel_data,
                                           school_name, semester,
                                           footer_text, date_text)
    cur = pb
    for el in elements:
        cur.addnext(el)
        cur = el
    return cur


TYPE_LABELS = {'class': '按班级', 'teacher': '按教师', 'groups': '分组信息'}


VALID_VIEW_TYPES = {'class', 'teacher', 'groups'}


@api_view(['POST'])
def export_word(request):
    """导出课表为 Word (.docx) 或 ZIP

    请求体::

        {
            "result_id": 1,
            "view_types": ["class", "teacher", "groups"],  // 多选
            "merge": true,                                 // 是否合并到一个文件
            "school_name": "某某某学校",
            "semester": "2025年下学期",
            "footer": "教导处",
            "date": "2025年8月"
        }
    """
    result_id = request.data.get('result_id')
    school_name = request.data.get('school_name', '').strip()
    semester = request.data.get('semester', '').strip()
    footer_text = request.data.get('footer', '').strip()
    date_text = request.data.get('date', '').strip()

    if not result_id:
        return Response({'error': '缺少 result_id'}, status=status.HTTP_400_BAD_REQUEST)

    # 支持多选 view_types，也兼容旧单值 view_type
    view_types = request.data.get('view_types')
    if not view_types:
        view_type = request.data.get('view_type', 'class')
        if view_type not in VALID_VIEW_TYPES:
            return Response({'error': f'无效的 view_type: {view_type}'},
                            status=status.HTTP_400_BAD_REQUEST)
        view_types = [view_type]

    for vt in view_types:
        if vt not in VALID_VIEW_TYPES:
            return Response({'error': f'无效的 view_type: {vt}'},
                            status=status.HTTP_400_BAD_REQUEST)

    try:
        ScheduleResult.objects.get(pk=result_id)
    except ScheduleResult.DoesNotExist:
        return Response({'error': '排课结果不存在'}, status=status.HTTP_404_NOT_FOUND)

    merge = request.data.get('merge', False)
    timetable_types = [vt for vt in view_types if vt in ('class', 'teacher')]
    has_groups = 'groups' in view_types

    if merge or len(view_types) == 1:
        # ── 合并为单个 docx ──
        page_items = []
        for vt in timetable_types:
            if vt == 'class':
                targets = list(SchoolClass.objects.all().order_by('grade', 'name'))
            else:
                targets = list(Teacher.objects.all().order_by('id'))
            if targets:
                cell_data_maps = _build_entry_maps(result_id, vt, targets)
                for t in targets:
                    page_items.append((t, cell_data_maps.get(t.id, [])))

        if not page_items and not has_groups:
            return Response({'error': '没有可导出的对象'},
                            status=status.HTTP_400_BAD_REQUEST)

        if not page_items and has_groups:
            # 仅导出分组信息
            docx_bytes = _generate_groups_docx_bytes(
                result_id, school_name, semester, footer_text, date_text
            )
            file_label = f'分组信息_{result_id}'
        else:
            # 先生成课表页
            doc = Document(str(TEMPLATE_PATH))
            body = doc.element.body

            sect_pr = body.find(qn('w:sectPr'))
            if sect_pr is not None:
                pg_mar = sect_pr.find(qn('w:pgMar'))
                if pg_mar is not None:
                    pg_mar.set(qn('w:bottom'), '450')
            all_children = list(body)

            title_el = all_children[IDX_TITLE]
            subtitle_el = all_children[IDX_SUBTITLE]
            table_el = all_children[IDX_TABLE]
            spacer_els = []
            for i in range(IDX_SPACER_START, IDX_SPACER_END + 1):
                if i < len(all_children) - 1:
                    spacer_els.append(all_children[i])
            footer_el_xml = all_children[IDX_FOOTER]
            date_el_xml = all_children[IDX_DATE]

            sect_pr = list(body)[-1]

            page_template = [title_el, subtitle_el, table_el] + spacer_els + [footer_el_xml, date_el_xml]

            # 填充第一页
            first_target, first_data = page_items[0]
            _set_subtitle_text(subtitle_el, f'{school_name}  {semester}', first_target.name)
            _fill_table(table_el, first_data)
            _set_footer_center(footer_el_xml, footer_text)
            _set_footer_center(date_el_xml, date_text)

            # 后续课表页
            last_page_end = date_el_xml
            for target, data in page_items[1:]:
                pb = _make_page_break()
                last_page_end.addnext(pb)
                cloned = _clone_page_elements(body, page_template)
                cl_title, cl_subtitle, cl_table = cloned[0], cloned[1], cloned[2]
                cl_footer, cl_date = cloned[-2], cloned[-1]

                _set_subtitle_text(cl_subtitle, f'{school_name}  {semester}', target.name)
                _fill_table(cl_table, data)
                _set_footer_center(cl_footer, footer_text)
                _set_footer_center(cl_date, date_text)

                last_page_end = _insert_sequence_after(pb, cloned)

            # 如果需要合并分组信息，追加到课表之后
            if has_groups:
                combined_data, travel_data = _get_groups_data(result_id)
                last_page_end = _append_groups_to_body(
                    body, combined_data, travel_data,
                    school_name, semester, footer_text, date_text,
                    last_page_end
                )

            # 确保 sectPr 在最后
            if list(body)[-1] is not sect_pr:
                body.remove(sect_pr)
                body.append(sect_pr)

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            docx_bytes = buf.getvalue()

            file_label = f'课表_{"_".join(view_types)}_{result_id}'

        disposition = f'attachment; filename="{file_label}.docx"'

        return HttpResponse(
            docx_bytes,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': disposition},
        )

    else:
        # ── 分开导出，打包为 zip ──
        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
            for vt in timetable_types:
                if vt == 'class':
                    targets = list(SchoolClass.objects.all().order_by('grade', 'name'))
                else:
                    targets = list(Teacher.objects.all().order_by('id'))

                if not targets:
                    continue

                docx_bytes = _build_single_docx_bytes(
                    result_id, vt, targets,
                    school_name, semester, footer_text, date_text
                )

                filename = f'课表_{TYPE_LABELS.get(vt, vt)}_{result_id}.docx'
                zf.writestr(filename, docx_bytes)

            if has_groups:
                docx_bytes = _generate_groups_docx_bytes(
                    result_id, school_name, semester, footer_text, date_text
                )
                filename = f'分组信息_{result_id}.docx'
                zf.writestr(filename, docx_bytes)

        zip_buf.seek(0)

        file_label = f'课表_{result_id}'
        disposition = f'attachment; filename="{file_label}.zip"'

        return HttpResponse(
            zip_buf.getvalue(),
            content_type='application/zip',
            headers={'Content-Disposition': disposition},
        )
